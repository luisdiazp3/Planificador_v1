import streamlit as st
from langchain_core.prompts import PromptTemplate
#from langchain_community.chat_models import ChatModel
from langchain_core.output_parsers import StrOutputParser
from langchain_openai import ChatOpenAI
from langchain_community.chat_models import ChatOpenAI
import os
import openai
import sys
from pydantic import BaseModel, ValidationError, field_validator
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from docx import Document

# Carga las variables de entorno desde el archivo .env
#load_dotenv()

# Obtiene la API key de las variables de entorno
#openai_api_key = os.getenv("OPENAI_API_KEY")
groq_api_key = st.secrets["GROQ_API_KEY"]
#GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# LOGO:
logo_path = "Guias/planifica.png"  # Ruta al archivo del logo
with st.sidebar:
    st.image(logo_path, width=200)  # Ajusta el ancho según sea necesario
logo_path = "Guias/planifica2.png"  # Ruta al archivo del logo
with st.sidebar:
    st.image(logo_path, width=200)  # Ajusta el ancho según sea necesario

# Función para generar un archivo Word
def generar_word(respuesta, nombre_archivo="respuesta.docx"):
    doc = Document()
    doc.add_heading("Propuesta de Sesión", level=1)
    doc.add_paragraph(respuesta)
    doc.save(nombre_archivo)
    return nombre_archivo

def obtenerRespuesta(nombre_sesion, area_curric, nivel, grado, competencia, desempeño, criterio_eval, enfoques, caracteristicas, materiales, tiempo_inicio, tiempo_desarro, tiempo_cierre):
    #llm = ChatOpenAI(openai_api_key=openai_api_key, model='gpt-3.5-turbo', temperature=0.2, max_tokens=2048)
    modelo = ChatGroq(groq_api_key=groq_api_key, model_name="llama3-70b-8192")

    template = """
        Responde en español latinoamericano. Actúa como un profesor del area de {area_curric} del nivel {nivel} y grado {grado} que desarrollará una propuesta de sesión
        de aprendizaje denominada {nombre_sesion} y espera observar el siguiente desempeño en los estudiantes: {desempeño}, que se alinea a la competencia: {competencia}.
        Desarrolla una sesión de aprendizaje que tenga la siguiente secuencia didactica inicio en {tiempo_inicio} minutos,
        el desarrollo en {tiempo_desarro} minutos y el cierre en {tiempo_cierre} minutos. 
        Considera en el inicio de la clase: un saludo inicial, una breve motivación, preguntas de recojo de saberes previos 
        y luego una situación que los motive al desarrollo de la clase, para el desarrollo formula actividades
        de trabajo en equipo, con mucha participación de los estudiantes, dispongo de los siguientes materiales: {materiales},
        y los estudiantes tienen las siguientes caracteristicas: {caracteristicas}. 
        Para el cierre considera preguntas para consolidar el tema que tenga el enfoque {enfoques}.
        Ahora elabora una rúbrica de evaluación, teniendo en cuenta los siguientes criterios: {criterio_eval}, 
        asegúrate de que la rúbrica incluya los siguientes niveles de desempeño: Aprendizaje en Inicio, Aprendizaje en Proceso, Aprendizaje Logrado y Aprendizaje Destacado. Presenta la rúbrica en formato de tabla.     

             """
    prompt = PromptTemplate(
        input_variables=["area_curric", "nivel", "grado", "nombre_sesion", "desempeño", "competencia", "tiempo_inicio", "tiempo_desarro", "tiempo_cierre", "materiales", "caracteristicas", "enfoques", "criterio_eval"],
        template=template
    )
        
    promt_formateado = prompt.format(area_curric=area_curric, nivel=nivel, grado=grado, nombre_sesion=nombre_sesion, desempeño=desempeño, competencia=competencia, tiempo_inicio=tiempo_inicio, tiempo_desarro=tiempo_desarro, tiempo_cierre=tiempo_cierre, materiales=materiales, caracteristicas=caracteristicas, enfoques=enfoques, criterio_eval=criterio_eval)

    chain = modelo | StrOutputParser()

    response = chain.invoke(promt_formateado)

    return response

# Interfaz de usuario
st.header("Planificador de Clases")
st.write("""Empieza a planificar la experiencia de aprendizaje, planteando una situación significativa que
          sea retadora para tus estudiantes, de acuerdo a sus características y necesidades de aprendizaje.
         Asimismo, selecciona las competencias y capacidades que desarrollarán los estudiantes para afrontar 
         el reto o los retos planteados en la situación. Para ello, debes contar con información diagnóstica
         del progreso de los aprendizajes de tus estudiantes.""")


nombre_sesion = st.text_input('Ingresa el nombre de la sesión de aprendizaje')
st.markdown(
    """
    <p style="font-size:14px; color:gray;">
        Para escribir el título, tener en cuenta la situación significativa planteada. 
    </p>
    """, unsafe_allow_html=True
)

# Línea divisoria
st.divider()

col2 = st.columns([10])[0]
with col2:
    nivel = st.selectbox(
        'Ingresa nivel',
        ('Elejir opción', 'Inicial', 'Primaria', 'Secundaria'),
        index=0  # Opción por defecto
    )

# Definir opciones dinámicas para el segundo selectbox
if nivel == "Primaria":
    opciones_grado = ('Elejir opción','Primero', 'Segundo', 'Tercero', 'Cuarto', 'Quinto', 'Sexto')
elif nivel == "Secundaria":
    opciones_grado = ('Elejir opción','Primero', 'Segundo', 'Tercero', 'Cuarto', 'Quinto')
else:  # Caso para Inicial u otro nivel
    opciones_grado = ('Elejir opción','3 años', '4 años', '5 años')

# Línea divisoria
st.divider()

col3 = st.columns([10])[0]
with col3:
    grado = st.selectbox(
        'Ingresa grado',
        opciones_grado,
        index=0  # Opción por defecto
    )

# Línea divisoria
st.divider()

nombre_sesion = st.text_area('Situación significativa')
st.markdown(
    """
    <p style="font-size:14px; color:gray;">
        Revisa el diagnóstico de estudiantes. Hazte las siguientes preguntas:
        ¿Qué les gustaría aprender a mis estudiantes?
        ¿Qué problemas pueden estar experimentando?
        ¿Cuáles son sus necesidades de aprendizaje? 
    </p>
    """, unsafe_allow_html=True
)

if nivel == "Primaria":
    opcion_area_curric = ('Elejir opción', 'Educación Física', 'Arte y Cultura',  'Comunicación', 
                   'Castellano como segunda lengua', 'Inglés como lengua extranjera',
                   'Matemática', 'Educación Religiosa', 'Transversal',  
                     )
elif nivel == "Secundaria":
    opcion_area_curric = ('Elejir opción', 'Educación Para el Trabajo', 'Educación Física', 'Arte y Cultura', 'Comunicación',
                   'Castellano como segunda lengua', 'Inglés como lengua extranjera', 'Matemática',)
else:  # Caso para Inicial u otro nivel
    opcion_area_curric = ('Elejir opción', 'Psicomotriz', 'Comunicación', 'Castellano como segunda lengua', 'Matemática',)

# Línea divisoria
st.divider()

col1 = st.columns([10])[0] 
with col1:
    area_curric = st.selectbox('Áreas curriculares', opcion_area_curric, index=0)

#col5 = st.columns([10])[0]
#with col5:                      
#    compt_transv = st.selectbox('Seleccione competencias transversales', ('Elejir opción', 'Se desenvuelve en los entornos virtuales generados por las TIC', 'Gestiona su aprendizaje de manera autónoma'),
#                                                  index=0)


if  area_curric == "Psicomotriz":
    opcion_competencia = ('Elejir opción', "Se desenvuelve de manera autónoma a través de su motricidad",)                     

elif  area_curric == "Comunicación":
    opcion_competencia = ('Elejir opción', "Se comunica oralmente en su lengua materna",
                            "Lee diversos tipos de textos escritos en lengua materna",
                            "Escribe diversos tipos de textos en su lengua materna",
                            "Crea proyectos desde los lenguajes artísticos",
                            "Se comunica oralmente en su lengua materna.",
                            "Lee diversos tipos de textos escritos en su lengua materna",
                            "Escribe diversos tipos de textos en lengua materna.",
                                                                                )

elif  area_curric == "Castellano como segunda lengua":
    opcion_competencia = ('Elejir opción', "Se comunica oralmente en castellano como segunda lengua.",
                        "Lee diversos tipos de textos escritos en castellano como segunda lengua.",
                        "Escribe diversos tipos de textos en castellano como segunda lengua.",
                                                                            )

elif  area_curric == 'Matemática':
    opcion_competencia = ('Elejir opción', "Resuelve problemas de cantidad.",
                        "Resuelve problemas de regularidad, equivalencia y cambio.",
                        "Resuelve problemas de gestión de datos e incertidumbre.",
                        "Resuelve problemas de forma, movimiento y localización.",  
                     )

elif  area_curric == 'Educación Física':
    opcion_competencia = ('Elejir opción', "Interactúa a través de sus habilidades sociomotrices",
                        "Se desenvuelve de manera autónoma a través de su motricidad",
                        "Asume una vida saludable",
                     )

elif  area_curric == 'Arte y Cultura':
    opcion_competencia = ('Elejir opción', "Crea proyectos desde los lenguajes artísticos",
                        "Aprecia de manera crítica manifestaciones artístico-culturales.",  
                     )

elif  area_curric == 'Inglés como lengua extranjera':
    opcion_competencia = ('Elejir opción', "Se comunica oralmente en inglés como lengua extranjera.",
                        "Lee diversos tipos de textos escritos en inglés como lengua extranjera.",
                        "Escribe diversos tipos de textos en inglés como lengua extranjera.",
                                         )

elif  area_curric == 'Educación Religiosa':
    opcion_competencia = ('Elejir opción', "Construye su identidad como persona humana, amada por Dios, digna, libre y trascendente, comprendiendo la doctrina de su propia religión, abierto al diálogo con las que le son cercanas",
                        "Asume la experiencia del encuentro personal y comunitario con Dios en su proyecto de vida en coherencia con su creencia religiosa",
                                     )

elif  area_curric == 'Educación Para el Trabajo':
    opcion_competencia = ('Elejir opción', 'Gestiona proyectos de emprendimiento económico o social',)


#elif  opcion_area_curric == 'Transversal':
#    opcion_competencia = ('Elejir opción', 'Se desenvuelve en los entornos virtuales generados por las TIC', 
#                          'Gestiona su aprendizaje de manera autónoma', 
#                     )

else:  # Caso para Transversal
    opcion_competencia = ('Elejir opción', 'Se desenvuelve en los entornos virtuales generados por las TIC', 
                          'Gestiona su aprendizaje de manera autónoma',)

# Línea divisoria
st.divider()

col4 = st.columns([10])[0] 
with col4:
    competencia = st.selectbox('Seleccione competencia', opcion_competencia, index=0)

# Línea divisoria
st.divider()

col6 = st.columns([10])[0]
with col6:                      # falta relacion de enfoques
    enfoques = st.selectbox('Seleccione enfoques', ('Elejir opción', "Enfoque de derechos",
"Enfoque Inclusivo o de Atención a la diversidad",
"Enfoque Intercultural",
"Enfoque Igualdad de Género",
"Enfoque Ambiental",
"Enfoque Orientación al bien común",
"Enfoque Búsqueda de la Excelencia"
), index=0)

# Línea divisoria
st.divider()

#submit = st.button("Generar el Formato de PerúEduca")

#st.write("""-------------------------------------------------------------------""")

st.write("""Para continuar con la creación de actividades para la Experiencia de Aprendizaje a través de la Inteligencia Artificial,
          proporciona los siguientes datos:""")

desempeño = st.text_area('Ingresa el desempeño a observar')    
# Nota informativa con enlace
st.markdown(
    """
    <p style="font-size:14px; color:gray;">
        Recuerda que los desempeños se encuentran en los programas curriculares. Sin embargo, 
        en algunas ocasiones, es necesario precisarlos para adaptarlos al contexto o a la situación significativa,
        siempre manteniendo sus niveles de exigencia. Si necesitas más información,  
        <a href="https://www.minedu.gob.pe/curriculo/" target="_blank">consúltala aquí.</a>.
    </p>
    """, unsafe_allow_html=True
)

# Línea divisoria
st.divider()


criterio_eval = st.text_area('Ingresa el criterio o criterios de evaluación')  
st.markdown(
    """
    <p style="font-size:14px; color:gray;">
        Recuerda los criterios se formulan a partir de los estándares y desempeños de grado descritos en el CNEB.
         Asimismo, deben incluir todas las capacidades de la competencia y ajustarse a la situación o problema a enfrentar,
         pues están alineados entre sí y describen la actuación correspondiente a observar. Si necesitas más información,  
        <a href="https://www.gob.pe/institucion/minedu/normas-legales/541161-094-2020-minedu" target="_blank">consúltala aquí.</a>.
    </p>
    """, unsafe_allow_html=True
)

# Línea divisoria
st.divider()

caracteristicas = st.text_area('¿Podrías compartir algunas características importantes a tener en cuenta respecto a los estudiantes? (sugerido)')
st.markdown(
    """
    <p style="font-size:14px; color:gray;">
    Ejemplos:
    </p>
    <style>
    ul.custom-list {
        font-size: 14px; !important; /* Cambia este valor al tamaño deseado */
        color: gray;
        margin-left: 15px;
        line-height: 1.0;
    }
    </style>
<ul class="custom-list">
    <li>Los estudiantes tienen conocimientos básicos en matemáticas y lectura, pero algunos requieren refuerzos en comprensión lectora y operaciones básicas.</li>
    <li>La mayoría de los estudiantes están motivados por aprender habilidades prácticas que puedan aplicar en su vida diaria, como el uso de tecnologías o resolución de problemas cotidianos.</li>
    <li>Dos estudiantes tienen déficit de atención con hiperactividad, y requieren de actividades con mayor nivel de interactividad.</li>
</ul>

    """, unsafe_allow_html=True
)

# Línea divisoria
st.divider()


materiales = st.text_area('Bríndanos información de algunos elementos, materiales, equipamiento tecnológico u otros recursos que emplearás durante la experiencia de aprendizaje. (Sugerido)') 
st.markdown(
    """
    <p style="font-size:14px; color:gray;">
    Ejemplos:
    </p>
    <style>
    ul.custom-list {
        font-size: 14px; !important; /* Cambia este valor al tamaño deseado */
        color: gray;
        margin-left: 15px;
        line-height: 1.0;
    }
    </style>
<ul class="custom-list">
    <li>Para esta sesión, cuento con cartulinas, plumones y papelotes. Los estudiantes están ubicados en carpetas individuales, pero tienen la posibilidad de agruparse en equipos de hasta seis participantes.</li>
    <li>En esta sesión, trabajaré en el aula de innovación pedagógica. Cada estudiante dispondrá de una laptop con conexión a internet. Además, cuento con un proyector y parlantes. Las carpetas, sin embargo, no pueden ser movilizadas.</li>
</ul>

    """, unsafe_allow_html=True
)

# Línea divisoria
st.divider()


tiempo_inicio = st.number_input('¿Cuántos minutos durará el inicio de la clase?. Ingresa un número')

tiempo_desarro = st.number_input('¿Cuántos minutos durará el desarrollo de la clase?. Ingresa un número')

tiempo_cierre = st.number_input('¿Cuántos minutos durará el cierre de la clase?. Ingresa un número')

# Línea divisoria
st.divider()

submit = st.button("Generar en pantalla")


if submit:
    respuesta = obtenerRespuesta(nombre_sesion, area_curric, nivel, grado, competencia, desempeño, criterio_eval, enfoques, caracteristicas, materiales, tiempo_inicio, tiempo_desarro, tiempo_cierre)
    st.write(respuesta)


# Botón para generar archivo Word
download = st.button("Generar en Word")
if download:
    respuesta = obtenerRespuesta(nombre_sesion, area_curric, nivel, grado, competencia, desempeño, criterio_eval, enfoques, caracteristicas, materiales, tiempo_inicio, tiempo_desarro, tiempo_cierre)
    archivo_word = generar_word(respuesta)
    with open(archivo_word, "rb") as file:
        st.download_button(
            label="Descargar Word",
            data=file,
            file_name=archivo_word,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# Nota informativa para los usuarios
st.info("⚠️ Recuerda que los modelos de inteligencia artificial pueden cometer errores, revisa y complementa la información brindada.")
